import os
import re
import shutil
from pathlib import Path
from typing import Optional

from PyQt5.QtCore import Qt, QSize, QPoint
from PyQt5.QtGui import QColor, QBrush
from PyQt5.QtWidgets import (
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QFileDialog,
    QListWidget,
    QListWidgetItem,
    QTableWidget,
    QTableWidgetItem,
    QHeaderView,
    QAbstractItemView,
    QInputDialog,
    QApplication,
)
from qfluentwidgets import (
    CardWidget,
    SubtitleLabel,
    BodyLabel,
    StrongBodyLabel,
    LineEdit,
    PushButton,
    PrimaryPushButton,
    InfoBar,
    IconWidget,
    FluentIcon as FIF,
    ComboBox,
    ToolButton,
)

from core.config_manager import get_config_manager
from core.font_manager import FontManager
from core.protocol_schema import (
    FieldSpec,
    compute_byte_positions,
    compute_bit_positions,
    compute_bit_group_info,
    load_csv,
    save_csv,
    generate_cpp_code,
    validate_fields,
)

TYPE_DISPLAY = [
    ("定值字节", "CONST"),
    ("不定字节", "ANY"),
    ("无符号1字节", "U8"),
    ("有符号1字节", "S8"),
    ("无符号2字节", "U16"),
    ("有符号2字节", "S16"),
    ("无符号4字节", "U32"),
    ("有符号4字节", "S32"),
    ("无符号1字节定标", "U8F"),
    ("有符号1字节定标", "S8F"),
    ("无符号2字节定标", "U16F"),
    ("有符号2字节定标", "S16F"),
    ("无符号4字节定标", "U32F"),
    ("有符号4字节定标", "S32F"),
    ("单精度浮点", "F32"),
    ("双精度浮点", "F64"),
    ("位字段", "BIT"),
]

class ProtocolEditorInterface(QWidget):
    """协议结构编辑器"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("protocolEditorInterface")
        self.config_manager = get_config_manager()
        self.current_dir = self.config_manager.get("protocol_csv_dir", "")
        self.current_file = None
        self._init_ui()
        if self.current_dir:
            self._load_folder(self.current_dir)

    def _init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(30, 20, 30, 20)
        layout.setSpacing(20)

        title = SubtitleLabel("协议建模工作台")
        title.setStyleSheet(f"color: #2D3748; font-size: {FontManager.get_font_size('large_title')}px;")
        layout.addWidget(title)

        desc = BodyLabel("编辑协议结构，生成 pack/unpack 与序列化 Schema 的 C++ 代码。")
        desc.setStyleSheet(f"color: #5A6A7A; font-size: {FontManager.get_font_size('body')}px;")
        layout.addWidget(desc)

        body_layout = QHBoxLayout()
        body_layout.setSpacing(20)
        body_layout.addWidget(self._create_folder_card(), 1)
        body_layout.addWidget(self._create_editor_card(), 2)
        layout.addLayout(body_layout, 1)

    def _create_folder_card(self):
        card = CardWidget()
        card.setStyleSheet("""
            CardWidget {
                background-color: rgba(255, 255, 255, 0.9);
                border-radius: 12px;
                border: 1px solid rgba(0, 0, 0, 0.06);
            }
        """)
        layout = QVBoxLayout(card)
        layout.setContentsMargins(24, 18, 24, 18)
        layout.setSpacing(12)

        header_layout = QHBoxLayout()
        header_layout.setSpacing(10)
        icon = IconWidget(FIF.FOLDER)
        icon.setFixedSize(28, 28)
        header_layout.addWidget(icon)

        title = StrongBodyLabel("协议资产库")
        title.setStyleSheet(f"font-size: {FontManager.get_font_size('title')}px; color: #2D3748;")
        header_layout.addWidget(title)
        header_layout.addStretch()
        layout.addLayout(header_layout)

        row_layout = QHBoxLayout()
        row_layout.setSpacing(10)

        self.folder_edit = LineEdit()
        self.folder_edit.setPlaceholderText("选择存放协议文件的文件夹")
        self.folder_edit.setFixedHeight(36)
        row_layout.addWidget(self.folder_edit, 1)

        browse_btn = ToolButton(FIF.FOLDER, self)
        browse_btn.setFixedSize(36, 36)
        browse_btn.setIconSize(QSize(20, 20))
        browse_btn.setToolTip("选择文件夹")
        browse_btn.clicked.connect(self._browse_folder)
        row_layout.addWidget(browse_btn)

        refresh_btn = ToolButton(FIF.SYNC, self)
        refresh_btn.setFixedSize(36, 36)
        refresh_btn.setIconSize(QSize(20, 20))
        refresh_btn.setToolTip("刷新列表")
        refresh_btn.clicked.connect(self._refresh_list)
        row_layout.addWidget(refresh_btn)

        layout.addLayout(row_layout)

        list_layout = QHBoxLayout()
        list_layout.setSpacing(12)

        self.csv_list = QListWidget()
        self.csv_list.setMinimumHeight(120)
        self.csv_list.itemSelectionChanged.connect(self._on_file_selected)
        list_layout.addWidget(self.csv_list, 1)

        btn_layout = QVBoxLayout()
        btn_layout.setSpacing(8)

        new_btn = PushButton("新建")
        new_btn.setFixedHeight(32)
        new_btn.clicked.connect(self._new_csv)
        btn_layout.addWidget(new_btn)

        copy_btn = PushButton("复制")
        copy_btn.setFixedHeight(32)
        copy_btn.clicked.connect(self._copy_csv)
        btn_layout.addWidget(copy_btn)

        rename_btn = PushButton("更名")
        rename_btn.setFixedHeight(32)
        rename_btn.clicked.connect(self._rename_csv)
        btn_layout.addWidget(rename_btn)

        delete_btn = PushButton("删除")
        delete_btn.setFixedHeight(32)
        delete_btn.clicked.connect(self._delete_csv)
        btn_layout.addWidget(delete_btn)

        btn_layout.addStretch()
        list_layout.addLayout(btn_layout)

        layout.addLayout(list_layout)
        return card

    def _create_editor_card(self):
        card = CardWidget()
        card.setStyleSheet("""
            CardWidget {
                background-color: rgba(255, 255, 255, 0.9);
                border-radius: 12px;
                border: 1px solid rgba(0, 0, 0, 0.06);
            }
        """)
        layout = QVBoxLayout(card)
        layout.setContentsMargins(24, 18, 24, 18)
        layout.setSpacing(12)

        header_layout = QHBoxLayout()
        header_layout.setSpacing(10)
        icon = IconWidget(FIF.EDIT)
        icon.setFixedSize(28, 28)
        header_layout.addWidget(icon)
        title = StrongBodyLabel("帧结构设计器")
        title.setStyleSheet(f"font-size: {FontManager.get_font_size('title')}px; color: #2D3748;")
        header_layout.addWidget(title)
        header_layout.addStretch()
        layout.addLayout(header_layout)

        self.table = QTableWidget(0, 8)
        self.table.setHorizontalHeaderLabels([
            "序号", "长度", "类型", "中文名", "英文名", "LSB", "默认值", "有效"
        ])
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(3, QHeaderView.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(4, QHeaderView.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(5, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(6, QHeaderView.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(7, QHeaderView.ResizeToContents)
        self.table.verticalHeader().setVisible(False)
        self.table.setColumnWidth(2, 160)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.SingleSelection)
        self.table.setEditTriggers(QAbstractItemView.DoubleClicked | QAbstractItemView.EditKeyPressed)
        self.table.itemChanged.connect(self._on_item_changed)
        layout.addWidget(self.table, 1)

        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(10)

        add_btn = PushButton("添加行")
        add_btn.setFixedHeight(32)
        add_btn.clicked.connect(self._add_row)
        btn_layout.addWidget(add_btn)

        remove_btn = PushButton("删除行")
        remove_btn.setFixedHeight(32)
        remove_btn.clicked.connect(self._remove_row)
        btn_layout.addWidget(remove_btn)

        copy_btn = PushButton("复制行")
        copy_btn.setFixedHeight(32)
        copy_btn.clicked.connect(self._copy_row)
        btn_layout.addWidget(copy_btn)

        move_up_btn = PushButton("上移")
        move_up_btn.setFixedHeight(32)
        move_up_btn.clicked.connect(self._move_row_up)
        btn_layout.addWidget(move_up_btn)

        move_down_btn = PushButton("下移")
        move_down_btn.setFixedHeight(32)
        move_down_btn.clicked.connect(self._move_row_down)
        btn_layout.addWidget(move_down_btn)

        btn_layout.addStretch()

        save_btn = PrimaryPushButton("保存")
        save_btn.setFixedHeight(32)
        save_btn.clicked.connect(self._save_csv)
        btn_layout.addWidget(save_btn)

        save_as_btn = PushButton("另存为")
        save_as_btn.setFixedHeight(32)
        save_as_btn.clicked.connect(self._save_as_csv)
        btn_layout.addWidget(save_as_btn)

        gen_btn = PrimaryPushButton("生成代码")
        gen_btn.setFixedHeight(32)
        gen_btn.clicked.connect(self._generate_code)
        btn_layout.addWidget(gen_btn)

        export_btn = PushButton("导出 Word")
        export_btn.setFixedHeight(32)
        export_btn.clicked.connect(self._export_word)
        btn_layout.addWidget(export_btn)

        layout.addLayout(btn_layout)
        return card

    def _browse_folder(self):
        directory = QFileDialog.getExistingDirectory(
            self, "选择协议文件夹", self.current_dir or os.path.expanduser("~")
        )
        if directory:
            self._load_folder(directory)

    def _load_folder(self, directory):
        self.current_dir = directory
        self.folder_edit.setText(directory)
        self.config_manager.set("protocol_csv_dir", directory)
        self._refresh_list()

    def _refresh_list(self):
        self.csv_list.clear()
        directory = self.folder_edit.text().strip()
        if not directory or not os.path.isdir(directory):
            return
        schema_files = []
        header_files = []
        for filename in sorted(os.listdir(directory)):
            lower = filename.lower()
            if lower.endswith(".csv"):
                schema_files.append(filename)
            elif lower.endswith(".h"):
                header_files.append(filename)
        for filename in schema_files + header_files:
            base_name, ext = os.path.splitext(filename)
            if ext.lower() == ".h":
                display_name = f"（代码）{base_name}"
                file_kind = "header"
            else:
                display_name = base_name
                file_kind = "schema"
            item = QListWidgetItem(display_name)
            item.setData(Qt.UserRole, os.path.join(directory, filename))
            item.setData(Qt.UserRole + 1, file_kind)
            self.csv_list.addItem(item)

    def _on_file_selected(self):
        items = self.csv_list.selectedItems()
        if not items:
            return
        item = items[0]
        path = item.data(Qt.UserRole)
        file_kind = item.data(Qt.UserRole + 1)
        if not path:
            return
        if file_kind == "header" or path.lower().endswith(".h"):
            try:
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read()
            except OSError:
                InfoBar.warning("提示", "读取代码失败", duration=2000, parent=self.window())
                return
            QApplication.clipboard().setText(content)
            display_name = os.path.splitext(os.path.basename(path))[0]
            InfoBar.success("已复制", f"已复制代码: {display_name}", duration=2000, parent=self.window())
            return
        self._load_csv_file(path)

    def _load_csv_file(self, path):
        self.current_file = path
        fields = load_csv(path)
        self._populate_table(fields)
        display_name = os.path.splitext(os.path.basename(path))[0]
        InfoBar.success("已加载", f"加载文件: {display_name}", duration=2000, parent=self.window())

    def _populate_table(self, fields):
        self.table.setRowCount(0)
        for field in fields:
            self._append_row(field)
        self._sync_index()

    def _append_row(self, field: Optional[FieldSpec] = None):
        row = self.table.rowCount()
        self.table.insertRow(row)

        index_item = QTableWidgetItem("")
        index_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)
        self.table.setItem(row, 0, index_item)

        length_item = QTableWidgetItem(str(field.length) if field else "")
        self.table.setItem(row, 1, length_item)

        type_combo = ComboBox()
        for display, value in TYPE_DISPLAY:
            type_combo.addItem(display, userData=value)
        type_combo.setMinimumWidth(140)
        type_combo.currentIndexChanged.connect(self._on_type_changed_from_sender)
        type_combo.blockSignals(True)
        if field:
            index = type_combo.findData(field.field_type)
            if index >= 0:
                type_combo.setCurrentIndex(index)
            else:
                type_combo.setCurrentIndex(type_combo.findData("U8"))
        else:
            type_combo.setCurrentIndex(type_combo.findData("U8"))
        self.table.setCellWidget(row, 2, type_combo)
        type_combo.setProperty("row", row)
        type_combo.blockSignals(False)

        self.table.setItem(row, 3, QTableWidgetItem(field.name_cn if field else ""))
        self.table.setItem(row, 4, QTableWidgetItem(field.name_en if field else ""))
        self.table.setItem(row, 5, QTableWidgetItem("" if not field or field.lsb is None else str(field.lsb)))
        if field is None:
            default_text = "0"
        else:
            default_text = "" if field.default is None else str(field.default)
        self.table.setItem(row, 6, QTableWidgetItem(default_text))

        valid_item = QTableWidgetItem()
        valid_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsUserCheckable | Qt.ItemIsSelectable)
        valid_item.setCheckState(Qt.Checked if field is None or field.is_valid else Qt.Unchecked)
        self.table.setItem(row, 7, valid_item)
        field_type = type_combo.currentData() if type_combo else "U8"
        self._set_valid_state_for_type(row, field_type)
        self._apply_row_style(row, field_type)
        if field is None:
            self._on_type_changed(row)

    def _sync_index(self):
        fields = self._collect_fields()
        positions = compute_byte_positions(fields)
        for row in range(self.table.rowCount()):
            item = self.table.item(row, 0)
            if not item:
                continue
            item.setText(positions[row] if row < len(positions) else "")
            type_combo = self.table.cellWidget(row, 2)
            if type_combo:
                type_combo.setProperty("row", row)
            field_type = type_combo.currentData() if type_combo else "U8"
            self._apply_row_style(row, field_type)

    def _row_from_widget(self, widget):
        if widget is None:
            return -1
        row = widget.property("row")
        if row is not None:
            try:
                return int(row)
            except (TypeError, ValueError):
                pass
        if widget.parent() is None:
            return -1
        pos = widget.mapTo(self.table.viewport(), QPoint(0, 0))
        return self.table.indexAt(pos).row()

    def _on_type_changed_from_sender(self, _idx):
        row = self._row_from_widget(self.sender())
        if row < 0:
            return
        self._on_type_changed(row)

    def _set_valid_state_for_type(self, row, field_type):
        valid_item = self.table.item(row, 7)
        if not valid_item:
            return
        if field_type == "BIT":
            valid_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable)
            valid_item.setCheckState(Qt.Checked)
            return
        valid_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsUserCheckable | Qt.ItemIsSelectable)

    def _apply_row_style(self, row, field_type):
        highlight = QBrush(QColor("#E6F7F7")) if field_type == "BIT" else QBrush()
        for col in range(self.table.columnCount()):
            item = self.table.item(row, col)
            if item:
                item.setBackground(highlight)
        type_combo = self.table.cellWidget(row, 2)
        if type_combo:
            if field_type == "BIT":
                type_combo.setStyleSheet("QComboBox { background-color: #E6F7F7; }")
            else:
                type_combo.setStyleSheet("")

    def _on_item_changed(self, item):
        if item.column() == 1:
            self._sync_index()

    def _on_type_changed(self, row):
        type_combo = self.table.cellWidget(row, 2)
        if not type_combo:
            return
        field_type = type_combo.currentData()
        if field_type == "BIT":
            self._set_valid_state_for_type(row, field_type)
            self._sync_index()
            return
        length_item = self.table.item(row, 1)
        if length_item is None:
            length_item = QTableWidgetItem("")
            self.table.setItem(row, 1, length_item)
        length_item.setText(str(self._default_length_for_type(field_type)))
        self._set_valid_state_for_type(row, field_type)
        self._sync_index()

    def _default_length_for_type(self, field_type):
        if field_type in ("CONST", "ANY", "U8", "S8", "U8F", "S8F"):
            return 1
        if field_type in ("U16", "S16", "U16F", "S16F"):
            return 2
        if field_type in ("U32", "S32", "U32F", "S32F", "F32"):
            return 4
        if field_type in ("F64",):
            return 8
        return 0

    def _add_row(self):
        self._append_row()
        self._sync_index()

    def _move_row_up(self):
        row = self.table.currentRow()
        if row <= 0:
            return
        fields = self._collect_fields()
        fields[row - 1], fields[row] = fields[row], fields[row - 1]
        self._populate_table(fields)
        self.table.setCurrentCell(row - 1, 0)

    def _move_row_down(self):
        row = self.table.currentRow()
        if row < 0 or row >= self.table.rowCount() - 1:
            return
        fields = self._collect_fields()
        fields[row + 1], fields[row] = fields[row], fields[row + 1]
        self._populate_table(fields)
        self.table.setCurrentCell(row + 1, 0)

    def _remove_row(self):
        row = self.table.currentRow()
        if row >= 0:
            self.table.removeRow(row)
            self._sync_index()

    def _increment_numbers(self, text):
        if not text:
            return text

        def repl(match):
            value = match.group(0)
            new_value = str(int(value) + 1)
            if value.startswith("0") and len(new_value) < len(value):
                return new_value.zfill(len(value))
            return new_value

        return re.sub(r"\d+", repl, text)

    def _copy_row(self):
        row = self.table.currentRow()
        if row < 0:
            return
        fields = self._collect_fields()
        if row >= len(fields):
            return
        source = fields[row]
        copied = FieldSpec(
            index=source.index,
            length=source.length,
            field_type=source.field_type,
            name_cn=self._increment_numbers(source.name_cn),
            name_en=self._increment_numbers(source.name_en),
            lsb=source.lsb,
            default=source.default,
            is_valid=source.is_valid,
        )
        fields.insert(row + 1, copied)
        self._populate_table(fields)
        self.table.setCurrentCell(row + 1, 0)

    def _collect_fields(self):
        fields = []
        for row in range(self.table.rowCount()):
            length_text = self.table.item(row, 1).text().strip() if self.table.item(row, 1) else ""
            try:
                length = int(length_text)
            except ValueError:
                length = 0
            type_combo = self.table.cellWidget(row, 2)
            field_type = type_combo.currentData() if type_combo else "U8"
            name_cn = self.table.item(row, 3).text().strip() if self.table.item(row, 3) else ""
            name_en = self.table.item(row, 4).text().strip() if self.table.item(row, 4) else ""
            lsb_text = self.table.item(row, 5).text().strip() if self.table.item(row, 5) else ""
            try:
                lsb = float(lsb_text) if lsb_text else None
            except ValueError:
                lsb = None
            default_text = self.table.item(row, 6).text().strip() if self.table.item(row, 6) else ""
            default_val = default_text if default_text else None
            valid_item = self.table.item(row, 7)
            is_valid = valid_item.checkState() == Qt.Checked if valid_item else False
            if field_type == "BIT":
                is_valid = True
            fields.append(FieldSpec(
                index=row + 1,
                length=length,
                field_type=field_type,
                name_cn=name_cn,
                name_en=name_en,
                lsb=lsb,
                default=default_val,
                is_valid=is_valid,
            ))
        return fields

    def _save_csv(self):
        if not self.current_file:
            self._save_as_csv()
            return
        fields = self._collect_fields()
        warnings = validate_fields(fields)
        if warnings:
            InfoBar.warning("提示", warnings[0], duration=2500, parent=self.window())
        try:
            save_csv(self.current_file, fields)
        except OSError as e:
            if getattr(e, "winerror", None) in (5, 32):
                InfoBar.error("保存失败", "文件被占用或无权限，请关闭占用程序或使用“另存为”", duration=3500, parent=self.window())
            else:
                InfoBar.error("保存失败", f"无法保存文件：{str(e)}", duration=3500, parent=self.window())
            return
        except Exception as e:
            InfoBar.error("保存失败", f"保存时发生异常：{str(e)}", duration=3500, parent=self.window())
            return

        display_name = os.path.splitext(os.path.basename(self.current_file))[0]
        InfoBar.success("保存成功", f"已保存: {display_name}", duration=2000, parent=self.window())
        self._refresh_list()

    def _save_as_csv(self):
        if not self.current_dir:
            InfoBar.warning("提示", "请先选择文件夹", duration=2000, parent=self.window())
            return
        name, ok = QInputDialog.getText(self, "另存为", "请输入文件名：")
        if not ok or not name:
            return
        filename = name if name.lower().endswith(".csv") else f"{name}.csv"
        path = os.path.join(self.current_dir, filename)
        self.current_file = path
        self._save_csv()

    def _new_csv(self):
        if not self.current_dir:
            InfoBar.warning("提示", "请先选择文件夹", duration=2000, parent=self.window())
            return
        name, ok = QInputDialog.getText(self, "新建", "请输入文件名：")
        if not ok or not name:
            return
        filename = name if name.lower().endswith(".csv") else f"{name}.csv"
        path = os.path.join(self.current_dir, filename)
        if os.path.exists(path):
            InfoBar.warning("提示", "文件已存在", duration=2000, parent=self.window())
            return
        save_csv(path, [])
        self._refresh_list()
        self._load_csv_file(path)

    def _copy_csv(self):
        items = self.csv_list.selectedItems()
        if not items:
            InfoBar.warning("提示", "请先选择协议文件", duration=2000, parent=self.window())
            return
        item = items[0]
        path = item.data(Qt.UserRole)
        file_kind = item.data(Qt.UserRole + 1)
        if not path or file_kind == "header" or path.lower().endswith(".h"):
            InfoBar.warning("提示", "只能复制协议文件", duration=2000, parent=self.window())
            return
        directory = os.path.dirname(path)
        base_name = os.path.splitext(os.path.basename(path))[0]
        suffix = "_copy"
        candidate = f"{base_name}{suffix}.csv"
        new_path = os.path.join(directory, candidate)
        counter = 2
        while os.path.exists(new_path):
            candidate = f"{base_name}{suffix}{counter}.csv"
            new_path = os.path.join(directory, candidate)
            counter += 1
        shutil.copy2(path, new_path)
        self._refresh_list()
        self._load_csv_file(new_path)

    def _rename_csv(self):
        items = self.csv_list.selectedItems()
        if not items:
            InfoBar.warning("提示", "请先选择协议或代码文件", duration=2000, parent=self.window())
            return
        item = items[0]
        path = item.data(Qt.UserRole)
        if not path:
            return
        directory = os.path.dirname(path)
        base_name, ext = os.path.splitext(os.path.basename(path))
        name, ok = QInputDialog.getText(self, "更名", "请输入新名称：", text=base_name)
        if not ok or not name:
            return
        new_base = os.path.splitext(name.strip())[0]
        if not new_base:
            return
        new_filename = f"{new_base}{ext}"
        new_path = os.path.join(directory, new_filename)
        if os.path.abspath(new_path) == os.path.abspath(path):
            return
        if os.path.exists(new_path):
            InfoBar.warning("提示", "文件已存在", duration=2000, parent=self.window())
            return
        try:
            os.rename(path, new_path)
        except OSError:
            InfoBar.warning("提示", "重命名失败", duration=2000, parent=self.window())
            return
        if self.current_file == path:
            self.current_file = new_path
        self.csv_list.blockSignals(True)
        self._refresh_list()
        for index in range(self.csv_list.count()):
            list_item = self.csv_list.item(index)
            if list_item.data(Qt.UserRole) == new_path:
                self.csv_list.setCurrentItem(list_item)
                break
        self.csv_list.blockSignals(False)
        InfoBar.success("重命名成功", f"已重命名: {new_base}", duration=2000, parent=self.window())

    def _delete_csv(self):
        items = self.csv_list.selectedItems()
        if not items:
            return
        item = items[0]
        path = item.data(Qt.UserRole)
        file_kind = item.data(Qt.UserRole + 1)
        if file_kind == "header" or (path and path.lower().endswith(".h")):
            InfoBar.warning("提示", "头文件不能在此删除。", duration=2000, parent=self.window())
            return
        if path and os.path.exists(path):
            os.remove(path)
            if self.current_file == path:
                self.current_file = None
                self.table.setRowCount(0)
            self._refresh_list()

    def _generate_code(self):
        if not self.current_file:
            InfoBar.warning("提示", "请先选择文件", duration=2000, parent=self.window())
            return
        fields = self._collect_fields()
        warnings = validate_fields(fields)
        if warnings:
            InfoBar.warning("提示", warnings[0], duration=2500, parent=self.window())
        frame_name = Path(self.current_file).stem
        code = generate_cpp_code(frame_name, fields)
        header_name = f"{frame_name}_protocol.h"
        output_path = os.path.join(self.current_dir, header_name)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(code)
        InfoBar.success("生成成功", f"已生成: {header_name}", duration=2500, parent=self.window())
        self._refresh_list()

    def _export_word(self):
        if not self.current_file:
            InfoBar.warning("提示", "请先选择协议文件", duration=2000, parent=self.window())
            return
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            InfoBar.warning("提示", "请先安装 python-docx", duration=2500, parent=self.window())
            return
        fields = self._collect_fields()
        if not fields:
            InfoBar.warning("提示", "没有可导出的数据", duration=2000, parent=self.window())
            return
        default_name = Path(self.current_file).stem + ".docx"
        default_path = os.path.join(self.current_dir or os.path.expanduser("~"), default_name)
        save_path, _ = QFileDialog.getSaveFileName(
            self, "保存 Word", default_path, "Word Document (*.docx)"
        )
        if not save_path:
            return
        if not save_path.lower().endswith(".docx"):
            save_path += ".docx"

        type_map = {value: display for display, value in TYPE_DISPLAY}
        byte_positions = compute_byte_positions(fields)
        bit_positions = compute_bit_positions(fields)
        group_by_index, group_info = compute_bit_group_info(fields)

        def format_byte_position(value):
            if not value:
                return ""
            return value.replace("B", "")

        def format_bit_position(value):
            if not value:
                return ""
            if "-" in value:
                return value
            return f"{value}-{value}"

        doc = Document()
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Microsoft YaHei"
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

        def set_cell_text(cell, text):
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        headers = ["序号", "字节序", "长度", "说明", "LSB", "默认值", "位序"]
        for idx, text in enumerate(headers):
            set_cell_text(table.cell(0, idx), text)

        group_rows: Dict[str, List[int]] = {}
        group_seq: Dict[str, int] = {}
        display_seq = 1
        for i, field in enumerate(fields):
            group_key = group_by_index.get(i)
            row = table.add_row().cells
            if group_key:
                group_rows.setdefault(group_key, []).append(len(table.rows) - 1)
            if group_key:
                if group_key not in group_seq:
                    group_seq[group_key] = display_seq
                    display_seq += 1
                set_cell_text(row[0], "")
            else:
                set_cell_text(row[0], str(display_seq))
                display_seq += 1
            byte_pos = format_byte_position(byte_positions[i] if i < len(byte_positions) else "")
            set_cell_text(row[1], byte_pos)
            if field.field_type == "BIT":
                length_text = f"{field.length}位" if field.length else ""
            else:
                length_value = field.length if field.length else self._default_length_for_type(field.field_type)
                length_text = f"{length_value}字节" if length_value else ""
            set_cell_text(row[2], length_text)
            type_text = type_map.get(field.field_type, field.field_type)
            desc = f"{field.name_cn or ''}，{type_text}" if (field.name_cn or type_text) else ""
            set_cell_text(row[3], desc)
            set_cell_text(row[4], "" if field.lsb is None else str(field.lsb))
            set_cell_text(row[5], "" if field.default is None else str(field.default))
            set_cell_text(row[6], format_bit_position(bit_positions[i] if (i < len(bit_positions) and field.field_type == "BIT") else ""))

            if group_key:
                info = group_info.get(group_key, {})
                indices = info.get("indices", [])
                if indices and i == indices[-1]:
                    total_bits = int(info.get("total_bits", 0))
                    container_bits = int(info.get("container_bits", 0))
                    remaining = container_bits - total_bits
                    if remaining > 0:
                        reserve_row = table.add_row().cells
                        group_rows[group_key].append(len(table.rows) - 1)
                        set_cell_text(reserve_row[0], "")
                        set_cell_text(reserve_row[1], byte_pos)
                        set_cell_text(reserve_row[2], f"{remaining}位")
                        reserve_desc = f"预留，{type_map.get('BIT', 'BIT')}"
                        set_cell_text(reserve_row[3], reserve_desc)
                        set_cell_text(reserve_row[4], "")
                        set_cell_text(reserve_row[5], "")
                        start_bit = total_bits + 1
                        end_bit = container_bits
                        set_cell_text(reserve_row[6], f"{start_bit}-{end_bit}")

        for group_key, rows in group_rows.items():
            if len(rows) <= 1:
                continue
            info = group_info.get(group_key, {})
            indices = info.get("indices", [])
            if not indices:
                continue
            seq_text = str(group_seq.get(group_key, indices[0] + 1))
            byte_text = format_byte_position(byte_positions[indices[0]] if indices[0] < len(byte_positions) else "")
            seq_cell = table.cell(rows[0], 0).merge(table.cell(rows[-1], 0))
            set_cell_text(seq_cell, seq_text)
            byte_cell = table.cell(rows[0], 1).merge(table.cell(rows[-1], 1))
            set_cell_text(byte_cell, byte_text)

        try:
            doc.save(save_path)
        except OSError:
            InfoBar.warning("提示", "保存失败，请检查路径或权限", duration=2500, parent=self.window())
            return
        InfoBar.success("导出成功", f"已保存: {os.path.basename(save_path)}", duration=2500, parent=self.window())
